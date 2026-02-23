"""
Servicio de generación de documentos Word (DOCX) para Fichas Técnicas.
Genera documentos que replican la estructura del template HTML/PDF existente.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import io
import base64


# ── Estilos y constantes ──────────────────────────────────────────────

FONT_NAME = "Arial"
FONT_SIZE_TITLE = Pt(11)
FONT_SIZE_SECTION_HEADER = Pt(7.5)
FONT_SIZE_LABEL = Pt(8)
FONT_SIZE_VALUE = Pt(8)
FONT_SIZE_SMALL = Pt(6.5)
FONT_SIZE_FOOTER = Pt(6)

COLOR_HEADER_BG = "E0E0E0"
COLOR_BORDER = "333333"
COLOR_ACCENT = "00A0B0"
COLOR_RED = "C41E3A"
COLOR_SELECTED_BG = "262626"


# ── Helpers de bajo nivel ─────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Configura los bordes de una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", COLOR_BORDER)}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=0, bottom=0, start=40, end=40):
    """Configura márgenes internos de celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{start}" w:type="dxa"/>'
        f'<w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_cell_width(cell, width_cm: float):
    """Fuerza el ancho exacto de una celda en cm."""
    width_dxa = int(width_cm * 567)  # 1 cm = 567 twips (dxa)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>'
    )
    # Remove existing tcW if any
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)


def _add_styled_run(paragraph, text, bold=False, size=None, color=None, font_name=FONT_NAME):
    """Agrega un run con estilo al párrafo."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _set_table_borders(table, color=COLOR_BORDER, size="4"):
    """Aplica bordes a toda la tabla."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _set_table_full_width(table, total_width_emu):
    """
    Force a table to occupy exactly `total_width_emu` (in EMU).
    Sets tblW type='dxa' so Word does not auto-shrink the table.
    """
    total_dxa = int(total_width_emu / 635)  # EMU → twips (1 twip = 635 EMU)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    # Remove existing tblW
    existing_w = tblPr.find(qn('w:tblW'))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>')
    tblPr.append(tblW)

    # Also set table layout to fixed so column widths are respected exactly
    existing_layout = tblPr.find(qn('w:tblLayout'))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    tblPr.append(layout)


def _remove_paragraph_spacing(paragraph):
    """Elimina espaciado antes y después del párrafo."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(9)


def _minimize_spacing(paragraph):
    """Reduce al mínimo el espaciado de un párrafo (para separadores entre tablas)."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = Pt(2)
    run = paragraph.add_run("")
    run.font.size = Pt(1)


def _decode_logo_base64(logo_b64: str) -> bytes:
    """Decodifica una imagen base64 data URI a bytes."""
    if not logo_b64:
        return None
    if "base64," in logo_b64:
        logo_b64 = logo_b64.split("base64,")[1]
    return base64.b64decode(logo_b64)


def _remove_cell_borders(cell):
    """Remove all borders from a cell (invisible)."""
    _set_cell_border(cell,
                     top={"val": "none", "sz": "0", "color": "FFFFFF"},
                     bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                     left={"val": "none", "sz": "0", "color": "FFFFFF"},
                     right={"val": "none", "sz": "0", "color": "FFFFFF"})


# ── Generación del documento ──────────────────────────────────────────

def generate_ficha_docx(ficha_data: dict, logo_left_b64: str = None, logo_right_b64: str = None) -> bytes:
    """
    Genera un documento Word (.docx) para una ficha técnica individual.
    """
    doc = Document()

    # ── Configurar página A4 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # Ancho útil en EMU (para pasar a _set_table_full_width)
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_cm = 19.0  # 21 - 1.0 - 1.0

    # ── Extraer datos ──
    ficha = ficha_data if isinstance(ficha_data, dict) else ficha_data
    servicio = ficha.get("servicio", {})
    tratamiento = ficha.get("tratamiento", {})
    productos = ficha.get("productos", [{}, {}, {}, {}])
    personal = ficha.get("personal_tecnico", ["", "", "", "", "", ""])
    obs_rec = ficha.get("obs_rec", {})

    # ── HEADER: Logo + Título + Número OS ──
    _add_header(doc, ficha, logo_left_b64, usable_width, usable_cm)

    spacer = doc.add_paragraph()
    _minimize_spacing(spacer)

    # ── INFO: Cliente, Fecha, Dirección, Distrito ──
    _add_info_row(doc, usable_width, usable_cm, [
        ("Cliente :", ficha.get("cliente", ""), 0.7),
        ("Fecha :", _format_date(ficha.get("fecha", "")), 0.3),
    ])
    _add_info_row(doc, usable_width, usable_cm, [
        ("Dirección :", ficha.get("direccion", ""), 0.7),
        ("Distrito :", ficha.get("distrito", ""), 0.3),
    ])

    # ── SERVICIO A EFECTUAR + DIAGNÓSTICO ──
    _add_section_service_diagnostic(doc, usable_width, usable_cm, servicio, ficha.get("diagnostico_area", ""))

    # ── CONDICIÓN SANITARIA ──
    _add_section_with_text(doc, usable_width, usable_cm, "CONDICIÓN SANITARIA DE LA ZONA CIRCUNDANTE",
                           ficha.get("condicion_sanitaria", ""), min_height=True)

    # ── TIPOS DE TRATAMIENTO ──
    _add_section_treatment(doc, usable_width, usable_cm, tratamiento)

    # ── PRODUCTOS QUÍMICOS Y/O BIOLÓGICOS ──
    _add_section_products(doc, usable_width, usable_cm, productos)

    # ── ACCIONES CORRECTIVAS ──
    _add_section_with_text(doc, usable_width, usable_cm, "ACCIONES CORRECTIVAS",
                           ficha.get("acciones_correctivas", ""))

    # ── ÁREAS TRATADAS ──
    _add_section_with_text(doc, usable_width, usable_cm, "ÁREAS TRATADAS",
                           ficha.get("areas_tratadas", ""))

    # ── PERSONAL TÉCNICO ──
    _add_section_personal(doc, usable_width, usable_cm, personal,
                          ficha.get("hora_inicio", ""),
                          ficha.get("hora_termino", ""),
                          ficha.get("numero_certificado", ""))

    # ── OBSERVACIONES Y RECOMENDACIONES ──
    _add_section_obs_rec(doc, usable_width, usable_cm, obs_rec)

    # ── EVALUACIÓN DE SATISFACCIÓN ──
    _add_section_satisfaction(doc, usable_width, usable_cm, ficha.get("satisfaccion", ""))

    # ── FIRMAS ──
    _add_signatures(doc, usable_width, usable_cm)

    # ── FOOTER ──
    _add_footer(doc, usable_width, usable_cm)

    # ── Reducir espaciado entre tablas ──
    _reduce_all_table_spacing(doc)

    # ── Generar bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ── Funciones auxiliares de secciones ──────────────────────────────────

def _reduce_all_table_spacing(doc):
    """Reduce el espaciado automático que Word agrega entre tablas."""
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            pf = paragraph.paragraph_format
            if pf.space_before is None:
                pf.space_before = Pt(0)
                pf.space_after = Pt(1)
                pf.line_spacing = Pt(2)
                for run in paragraph.runs:
                    if run.font.size is None:
                        run.font.size = Pt(1)


def _format_date(fecha_str: str) -> str:
    """Convierte fecha ISO (YYYY-MM-DD) a DD-MM-YYYY."""
    if not fecha_str:
        return ""
    try:
        parts = fecha_str.split(" ")[0].split("-")
        if len(parts) == 3:
            return f"{parts[2]}-{parts[1]}-{parts[0]}"
    except Exception:
        pass
    return fecha_str


def _add_header(doc, ficha, logo_left_b64, usable_width, usable_cm):
    """Header: Logo | Título | N° OS — forced full width."""
    # Column widths: logo 4.5cm, title remainder, OS 3.5cm
    col_logo = 4.5
    col_os = 3.5
    col_title = usable_cm - col_logo - col_os

    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(header_table, usable_width)

    header_table.columns[0].width = Cm(col_logo)
    header_table.columns[1].width = Cm(col_title)
    header_table.columns[2].width = Cm(col_os)

    # Force cell widths explicitly
    for row in header_table.rows:
        _set_cell_width(row.cells[0], col_logo)
        _set_cell_width(row.cells[1], col_title)
        _set_cell_width(row.cells[2], col_os)

    # Logo
    cell_logo = header_table.cell(0, 0)
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if logo_left_b64:
        try:
            logo_bytes = _decode_logo_base64(logo_left_b64)
            if logo_bytes:
                logo_stream = io.BytesIO(logo_bytes)
                p = cell_logo.paragraphs[0]
                run = p.add_run()
                run.add_picture(logo_stream, width=Cm(4.0))
        except Exception as e:
            print(f"[WordService] Error inserting logo: {e}")
    _remove_paragraph_spacing(cell_logo.paragraphs[0])

    # Título
    cell_title = header_table.cell(0, 1)
    cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_title = cell_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p_title, "FICHA TÉCNICA DE EVALUACIÓN\nDE ACTIVIDADES",
                    bold=True, size=FONT_SIZE_TITLE, color=COLOR_BORDER)
    _remove_paragraph_spacing(p_title)

    # N° OS
    cell_os = header_table.cell(0, 2)
    cell_os.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p_os = cell_os.paragraphs[0]
    p_os.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    os_numero = ficha.get("os_numero", "")
    os_display = os_numero.replace("OS-", "").replace("-", "") if os_numero else "00000"
    _add_styled_run(p_os, f"N° {os_display}", bold=True, size=Pt(11), color=COLOR_RED)
    p_os2 = cell_os.add_paragraph()
    p_os2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(p_os2, "O.S.N° ", bold=True, size=Pt(7))
    _add_styled_run(p_os2, "________", size=Pt(7))
    _remove_paragraph_spacing(p_os)
    _remove_paragraph_spacing(p_os2)

    # Remove borders
    for row in header_table.rows:
        for cell in row.cells:
            _remove_cell_borders(cell)


def _add_info_row(doc, usable_width, usable_cm, fields):
    """Agrega una fila de información (label: valor) con tabla invisible al 100% ancho."""
    num_cols = len(fields) * 2
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    col_idx = 0
    for label, value, ratio in fields:
        field_width = usable_cm * ratio
        label_width = 2.0
        value_width = field_width - label_width

        # Label column
        cell_label = table.cell(0, col_idx)
        table.columns[col_idx].width = Cm(label_width)
        _set_cell_width(cell_label, label_width)
        p = cell_label.paragraphs[0]
        _add_styled_run(p, label, bold=True, size=FONT_SIZE_LABEL)
        _remove_paragraph_spacing(p)
        _remove_cell_borders(cell_label)

        # Value column
        cell_value = table.cell(0, col_idx + 1)
        table.columns[col_idx + 1].width = Cm(value_width)
        _set_cell_width(cell_value, value_width)
        p2 = cell_value.paragraphs[0]
        _add_styled_run(p2, str(value) if value else "", size=FONT_SIZE_VALUE)
        _remove_paragraph_spacing(p2)
        _set_cell_border(cell_value,
                         top={"val": "none", "sz": "0", "color": "FFFFFF"},
                         bottom={"val": "single", "sz": "4", "color": COLOR_BORDER},
                         left={"val": "none", "sz": "0", "color": "FFFFFF"},
                         right={"val": "none", "sz": "0", "color": "FFFFFF"})

        col_idx += 2

    return table


def _add_section_header_row(table, row_idx, text, num_cols):
    """Agrega encabezado de sección gris en una tabla."""
    if num_cols > 1:
        cell = table.cell(row_idx, 0)
        for i in range(1, num_cols):
            cell = cell.merge(table.cell(row_idx, i))
    else:
        cell = table.cell(row_idx, 0)

    _set_cell_shading(cell, COLOR_HEADER_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, text, bold=True, size=FONT_SIZE_SECTION_HEADER)
    _remove_paragraph_spacing(p)


def _add_section_service_diagnostic(doc, usable_width, usable_cm, servicio, diagnostico):
    """Sección Servicio a Efectuar + Diagnóstico lado a lado."""
    # Service side ≈ 42%, diagnostic ≈ 58% (matching PDF flex:1 / flex:1 with checkbox col)
    col_svc_name = usable_cm * 0.35
    col_svc_chk = 1.0
    col_diag_sep = 0.3
    col_diag = usable_cm - col_svc_name - col_svc_chk - col_diag_sep

    table = doc.add_table(rows=6, cols=4)
    _set_table_borders(table, size="6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    table.columns[0].width = Cm(col_svc_name)
    table.columns[1].width = Cm(col_svc_chk)
    table.columns[2].width = Cm(col_diag_sep)
    table.columns[3].width = Cm(col_diag)

    # Force cell widths on every row
    for row in table.rows:
        _set_cell_width(row.cells[0], col_svc_name)
        _set_cell_width(row.cells[1], col_svc_chk)

    # Header izquierdo
    cell_srv_header = table.cell(0, 0).merge(table.cell(0, 1))
    _set_cell_shading(cell_srv_header, COLOR_HEADER_BG)
    p = cell_srv_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, "SERVICIO A EFECTUAR", bold=True, size=FONT_SIZE_SECTION_HEADER)
    _remove_paragraph_spacing(p)

    # Header derecho
    cell_diag_header = table.cell(0, 2).merge(table.cell(0, 3))
    _set_cell_shading(cell_diag_header, COLOR_HEADER_BG)
    p = cell_diag_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, "DIAGNÓSTICO DEL ÁREA A TRATAR", bold=True, size=FONT_SIZE_SECTION_HEADER)
    _remove_paragraph_spacing(p)

    # Servicios
    services = [
        ("1. DESINFECCIÓN", servicio.get("desinfeccion", False)),
        ("2. LIMPIEZA DE AMBIENTES", servicio.get("limpieza_ambientes", False)),
        ("3. LIMPIEZA DE POZOS SÉPTICOS", servicio.get("limpieza_pozos_septicos", False)),
        ("4. LIMPIEZA Y DESINFECCIÓN DE RESERVORIOS DE AGUA", servicio.get("limpieza_reservorios", False)),
    ]

    # Diagnóstico (merged rows 1-5, cols 2-3)
    diag_cell = table.cell(1, 2).merge(table.cell(5, 3))
    p_diag = diag_cell.paragraphs[0]
    _add_styled_run(p_diag, diagnostico or "", size=FONT_SIZE_VALUE)
    _remove_paragraph_spacing(p_diag)
    _set_cell_margins(diag_cell, top=30, bottom=30, start=50, end=50)

    for i, (label, checked) in enumerate(services):
        row_idx = i + 1
        cell_name = table.cell(row_idx, 0)
        p = cell_name.paragraphs[0]
        _add_styled_run(p, label, size=FONT_SIZE_SMALL)
        _remove_paragraph_spacing(p)
        _set_cell_margins(cell_name, top=15, bottom=15, start=40, end=20)

        cell_chk = table.cell(row_idx, 1)
        p = cell_chk.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(p, "☒" if checked else "☐", size=Pt(9))
        _remove_paragraph_spacing(p)

    # Fila vacía extra
    cell_empty = table.cell(5, 0).merge(table.cell(5, 1))
    _remove_paragraph_spacing(cell_empty.paragraphs[0])


def _add_section_with_text(doc, usable_width, usable_cm, title, text, min_height=False):
    """Agrega una sección con header gris y texto."""
    table = doc.add_table(rows=2, cols=1)
    _set_table_borders(table, size="6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    table.columns[0].width = Cm(usable_cm)

    cell_header = table.cell(0, 0)
    _set_cell_shading(cell_header, COLOR_HEADER_BG)
    p = cell_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, title, bold=True, size=FONT_SIZE_SECTION_HEADER)
    _remove_paragraph_spacing(p)

    cell_content = table.cell(1, 0)
    p = cell_content.paragraphs[0]
    _add_styled_run(p, text or "", size=FONT_SIZE_VALUE)
    _remove_paragraph_spacing(p)
    _set_cell_margins(cell_content, top=20, bottom=20, start=40, end=40)

    if min_height and not text:
        p.add_run("\n")


def _add_section_treatment(doc, usable_width, usable_cm, tratamiento):
    """Sección Tipos de Tratamiento."""
    # PDF: 25%, 10%, 30%, 25%, 10%
    c0 = usable_cm * 0.25
    c1 = usable_cm * 0.10
    c2 = usable_cm * 0.30
    c3 = usable_cm * 0.25
    c4 = usable_cm * 0.10

    table = doc.add_table(rows=4, cols=5)
    _set_table_borders(table, size="6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    table.columns[0].width = Cm(c0)
    table.columns[1].width = Cm(c1)
    table.columns[2].width = Cm(c2)
    table.columns[3].width = Cm(c3)
    table.columns[4].width = Cm(c4)

    _add_section_header_row(table, 0, "TIPOS DE TRATAMIENTO", 5)

    treatments = [
        [("Pulverizado", tratamiento.get("pulverizado", False)),
         ("Thermonebulizado", tratamiento.get("thermonebulizado", False))],
        [("Atomizado", tratamiento.get("atomizado", False)),
         ("Nebulizado ULV", tratamiento.get("nebulizado_ulv", False))],
        [("Otros: " + (tratamiento.get("otros", "") or ""), None), (None, None)],
    ]

    for row_i, row_data in enumerate(treatments):
        row_idx = row_i + 1

        label_l, checked_l = row_data[0]
        cell = table.cell(row_idx, 0)
        p = cell.paragraphs[0]
        _add_styled_run(p, label_l, size=FONT_SIZE_VALUE)
        _remove_paragraph_spacing(p)
        _set_cell_margins(cell, top=15, bottom=15, start=60, end=20)

        cell_chk = table.cell(row_idx, 1)
        p = cell_chk.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if checked_l is not None:
            _add_styled_run(p, "☒" if checked_l else "☐", size=Pt(9))
        _remove_paragraph_spacing(p)

        _remove_paragraph_spacing(table.cell(row_idx, 2).paragraphs[0])

        label_r, checked_r = row_data[1]
        cell = table.cell(row_idx, 3)
        p = cell.paragraphs[0]
        if label_r:
            _add_styled_run(p, label_r, size=FONT_SIZE_VALUE)
        _remove_paragraph_spacing(p)
        _set_cell_margins(cell, top=15, bottom=15, start=60, end=20)

        cell_chk = table.cell(row_idx, 4)
        p = cell_chk.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if checked_r is not None:
            _add_styled_run(p, "☒" if checked_r else "☐", size=Pt(9))
        _remove_paragraph_spacing(p)


def _add_section_products(doc, usable_width, usable_cm, productos):
    """Sección Productos Químicos y/o Biológicos."""
    num_rows = 1 + 1 + len(productos)
    table = doc.add_table(rows=num_rows, cols=7)
    _set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    # Column proportions from PDF
    proportions = [0.16, 0.16, 0.12, 0.16, 0.10, 0.16, 0.14]
    for col_i, prop in enumerate(proportions):
        table.columns[col_i].width = Cm(usable_cm * prop)

    _add_section_header_row(table, 0, "PRODUCTOS QUÍMICOS Y/O BIOLÓGICOS UTILIZADOS", 7)

    headers = ["PRODUCTO", "COMPOSICIÓN", "LOTE", "FECHA DE\nVENCIMIENTO", "UNIDAD", "CONCENTRACIÓN", "CANTIDAD"]
    for col_i, header_text in enumerate(headers):
        cell = table.cell(1, col_i)
        _set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(p, header_text, bold=True, size=FONT_SIZE_SMALL)
        _remove_paragraph_spacing(p)

    for prod_i, producto in enumerate(productos):
        row_idx = prod_i + 2
        if isinstance(producto, dict):
            values = [
                producto.get("producto", ""),
                producto.get("composicion", ""),
                producto.get("lote", ""),
                producto.get("fecha_vencimiento", ""),
                producto.get("unidad", ""),
                producto.get("concentracion", ""),
                _format_cantidad(producto.get("cantidad", "")),
            ]
        else:
            values = [""] * 7

        for col_i, val in enumerate(values):
            cell = table.cell(row_idx, col_i)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_styled_run(p, str(val) if val else "", size=FONT_SIZE_SMALL)
            _remove_paragraph_spacing(p)


def _format_cantidad(cantidad) -> str:
    """Formatea cantidad con 4 decimales si es numérico."""
    if not cantidad:
        return ""
    try:
        return "{:.4f}".format(float(cantidad))
    except (ValueError, TypeError):
        return str(cantidad)


def _add_section_personal(doc, usable_width, usable_cm, personal, hora_inicio, hora_termino, num_certificado):
    """Sección Personal Técnico + Horarios + Certificado."""
    table = doc.add_table(rows=5, cols=6)
    _set_table_borders(table, size="6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    # Each of 6 cols = usable/6
    col_w = usable_cm / 6.0
    for i in range(6):
        table.columns[i].width = Cm(col_w)

    _add_section_header_row(table, 0, "PERSONAL TÉCNICO", 6)

    for i in range(3):
        cell_l = table.cell(i + 1, 0).merge(table.cell(i + 1, 2))
        p = cell_l.paragraphs[0]
        val_l = personal[i] if i < len(personal) else ""
        _add_styled_run(p, str(val_l) if val_l else "", size=FONT_SIZE_SMALL)
        _remove_paragraph_spacing(p)
        _set_cell_margins(cell_l, top=15, bottom=15, start=40, end=40)

        cell_r = table.cell(i + 1, 3).merge(table.cell(i + 1, 5))
        p = cell_r.paragraphs[0]
        val_r = personal[i + 3] if (i + 3) < len(personal) else ""
        _add_styled_run(p, str(val_r) if val_r else "", size=FONT_SIZE_SMALL)
        _remove_paragraph_spacing(p)
        _set_cell_margins(cell_r, top=15, bottom=15, start=40, end=40)

    # Time row: 3 separate fields
    time_cell_1 = table.cell(4, 0).merge(table.cell(4, 1))
    p1 = time_cell_1.paragraphs[0]
    _add_styled_run(p1, "HORA INICIO : ", bold=True, size=FONT_SIZE_SMALL)
    _add_styled_run(p1, str(hora_inicio) if hora_inicio else "", size=FONT_SIZE_SMALL)
    _remove_paragraph_spacing(p1)
    _set_cell_margins(time_cell_1, top=20, bottom=20, start=40, end=40)

    time_cell_2 = table.cell(4, 2).merge(table.cell(4, 3))
    p2 = time_cell_2.paragraphs[0]
    _add_styled_run(p2, "HORA TÉRMINO : ", bold=True, size=FONT_SIZE_SMALL)
    _add_styled_run(p2, str(hora_termino) if hora_termino else "", size=FONT_SIZE_SMALL)
    _remove_paragraph_spacing(p2)
    _set_cell_margins(time_cell_2, top=20, bottom=20, start=40, end=40)

    time_cell_3 = table.cell(4, 4).merge(table.cell(4, 5))
    p3 = time_cell_3.paragraphs[0]
    _add_styled_run(p3, "N° CERTIFICADO : ", bold=True, size=FONT_SIZE_SMALL)
    _add_styled_run(p3, str(num_certificado) if num_certificado else "", size=FONT_SIZE_SMALL)
    _remove_paragraph_spacing(p3)
    _set_cell_margins(time_cell_3, top=20, bottom=20, start=40, end=40)


def _add_section_obs_rec(doc, usable_width, usable_cm, obs_rec):
    """Sección Observaciones y Recomendaciones."""
    table = doc.add_table(rows=4, cols=2)
    _set_table_borders(table, size="6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    half = usable_cm / 2.0
    table.columns[0].width = Cm(half)
    table.columns[1].width = Cm(half)

    cell_obs = table.cell(0, 0)
    _set_cell_shading(cell_obs, COLOR_HEADER_BG)
    p = cell_obs.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, "OBSERVACIONES", bold=True, size=FONT_SIZE_SECTION_HEADER)
    _remove_paragraph_spacing(p)

    cell_rec = table.cell(0, 1)
    _set_cell_shading(cell_rec, COLOR_HEADER_BG)
    p = cell_rec.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, "RECOMENDACIONES", bold=True, size=FONT_SIZE_SECTION_HEADER)
    _remove_paragraph_spacing(p)

    obs_items = [
        obs_rec.get("observacion_a", ""),
        obs_rec.get("observacion_b", ""),
        obs_rec.get("observacion_c", ""),
    ]
    rec_items = [
        obs_rec.get("recomendacion_a", ""),
        obs_rec.get("recomendacion_b", ""),
        obs_rec.get("recomendacion_c", ""),
    ]
    letters = ["a)", "b)", "c)"]

    for i in range(3):
        cell = table.cell(i + 1, 0)
        p = cell.paragraphs[0]
        _add_styled_run(p, f"{letters[i]} ", bold=True, size=FONT_SIZE_VALUE)
        _add_styled_run(p, obs_items[i] or "", size=FONT_SIZE_VALUE)
        _remove_paragraph_spacing(p)
        _set_cell_margins(cell, top=10, bottom=10, start=30, end=30)
        _set_cell_border(cell, bottom={"val": "single", "sz": "2", "color": COLOR_BORDER})

        cell = table.cell(i + 1, 1)
        p = cell.paragraphs[0]
        _add_styled_run(p, f"{letters[i]} ", bold=True, size=FONT_SIZE_VALUE)
        _add_styled_run(p, rec_items[i] or "", size=FONT_SIZE_VALUE)
        _remove_paragraph_spacing(p)
        _set_cell_margins(cell, top=10, bottom=10, start=30, end=30)
        _set_cell_border(cell, bottom={"val": "single", "sz": "2", "color": COLOR_BORDER})


def _add_section_satisfaction(doc, usable_width, usable_cm, satisfaccion):
    """Sección Evaluación de Satisfacción del Cliente."""
    table = doc.add_table(rows=2, cols=1)
    _set_table_borders(table, size="6")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    table.columns[0].width = Cm(usable_cm)

    cell_header = table.cell(0, 0)
    _set_cell_shading(cell_header, COLOR_HEADER_BG)
    p = cell_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, "EVALUACIÓN DE SATISFACCIÓN DEL CLIENTE", bold=True, size=FONT_SIZE_SECTION_HEADER)
    _remove_paragraph_spacing(p)

    cell_content = table.cell(1, 0)
    p = cell_content.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_margins(cell_content, top=40, bottom=40, start=30, end=30)

    options = [
        ("😊", "Muy Satisfecho", "muy_satisfecho"),
        ("🙂", "Satisfecho", "satisfecho"),
        ("😐", "Regular", "regular"),
        ("☹️", "Insatisfecho", "insatisfecho"),
    ]

    for idx, (emoji_icon, label, key) in enumerate(options):
        is_selected = satisfaccion == key
        separator = "     " if idx < len(options) - 1 else ""

        if is_selected:
            _add_styled_run(p, f" {emoji_icon} ", size=FONT_SIZE_VALUE)
            run = _add_styled_run(p, f" {label} ", bold=True, size=FONT_SIZE_VALUE)
            run.font.color.rgb = RGBColor.from_string(COLOR_ACCENT)
        else:
            _add_styled_run(p, f" {emoji_icon} ", size=FONT_SIZE_VALUE)
            _add_styled_run(p, f" {label} ", size=FONT_SIZE_VALUE)

        if separator:
            _add_styled_run(p, separator, size=FONT_SIZE_VALUE)

    _remove_paragraph_spacing(p)


def _add_signatures(doc, usable_width, usable_cm):
    """Agrega sección de firmas."""
    for _ in range(3):
        spacer = doc.add_paragraph()
        _minimize_spacing(spacer)

    sig_spacer = doc.add_paragraph()
    pf = sig_spacer.paragraph_format
    pf.space_before = Cm(1.5)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(2)

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table, usable_width)

    col_w = usable_cm / 3.0
    for i in range(3):
        table.columns[i].width = Cm(col_w)

    labels = ["Responsable de Servicio", "Cliente", "Director Técnico"]
    for col_i, label in enumerate(labels):
        cell_line = table.cell(0, col_i)
        p = cell_line.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(p, "________________________", size=FONT_SIZE_SMALL)
        _remove_paragraph_spacing(p)

        cell_label = table.cell(1, col_i)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(p, label, size=FONT_SIZE_SMALL)
        _remove_paragraph_spacing(p)

        for cell in [cell_line, cell_label]:
            _remove_cell_borders(cell)


def _add_footer(doc, usable_width, usable_cm):
    """Agrega pie de página con información de contacto."""
    spacer = doc.add_paragraph()
    _minimize_spacing(spacer)

    footer_table = doc.add_table(rows=1, cols=3)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(footer_table, usable_width)

    col_w = usable_cm / 3.0
    for i in range(3):
        footer_table.columns[i].width = Cm(col_w)

    cell_left = footer_table.cell(0, 0)
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_styled_run(p, "📍 Mz J1 lote 20. Urb. Los Precursores. Surco. Lima", size=FONT_SIZE_FOOTER, color=COLOR_ACCENT)
    _remove_paragraph_spacing(p)

    cell_center = footer_table.cell(0, 1)
    p_email = cell_center.paragraphs[0]
    p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p_email, "✉ operaciones@hidroserviciosaa.com.pe", size=FONT_SIZE_FOOTER, color=COLOR_ACCENT)
    _remove_paragraph_spacing(p_email)
    p_phone = cell_center.add_paragraph()
    p_phone.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p_phone, "📞 +51 946 803 367", size=FONT_SIZE_FOOTER, color=COLOR_ACCENT)
    _remove_paragraph_spacing(p_phone)

    cell_right = footer_table.cell(0, 2)
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(p, "🌐 www.hidroserviciosaa.com.pe", size=FONT_SIZE_FOOTER, color=COLOR_ACCENT)
    _remove_paragraph_spacing(p)

    for row in footer_table.rows:
        for cell in row.cells:
            _remove_cell_borders(cell)
