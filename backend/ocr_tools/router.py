import mimetypes
import os
import re
import tempfile
import json
from pathlib import Path
from typing import Any, Dict, Literal, Tuple

import httpx  
from docx import Document  
from fastapi import APIRouter, BackgroundTasks, File, Form, HTTPException, UploadFile  
from fastapi.responses import FileResponse, JSONResponse  

from config import settings  
from .service import OCRConfigurationError, OCRServiceError, OllamaOCRService, create_ocr_service

router = APIRouter(prefix="/api/tools", tags=["ocr-tools"])

_ocr_service = create_ocr_service()

_OLLAMA_URL_MAX_LENGTH = 300
_OLLAMA_MODEL_MAX_LENGTH = 120


def _resolve_service(ollama_url: str, ollama_model: str) -> Any:
    """Return an on-the-fly OllamaOCRService if the caller supplies URL+model,
    otherwise fall back to the module-level singleton configured via env vars."""
    url = (ollama_url or "").strip()
    model = (ollama_model or "").strip()

    if not url or not model:
        return _ocr_service

    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="ollama_url debe comenzar con http:// o https://")
    if len(url) > _OLLAMA_URL_MAX_LENGTH or len(model) > _OLLAMA_MODEL_MAX_LENGTH:
        raise HTTPException(status_code=400, detail="ollama_url o ollama_model demasiado largo")

    try:
        return OllamaOCRService(model=model, base_url=url, backend_name="ollama-custom")
    except OCRConfigurationError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

_SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".bmp",
    ".webp",
    ".heic",
    ".heif",
}

_SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/tiff",
    "image/bmp",
    "image/webp",
    "image/heic",
    "image/heif",
}

_STRUCTURED_PRESETS: Dict[str, Dict[str, Any]] = {
    "general": {
        "schema_name": "documento_general",
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "titulo": {"type": "string"},
                "fecha_principal": {"type": "string"},
                "resumen": {"type": "string"},
                "entidades_clave": {"type": "array", "items": {"type": "string"}},
                "valores_clave": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "campo": {"type": "string"},
                            "valor": {"type": "string"},
                        },
                        "required": ["campo", "valor"],
                    },
                },
            },
            "required": ["titulo", "fecha_principal", "resumen", "entidades_clave", "valores_clave"],
        },
        "prompt": "Extrae los datos clave del documento y rellena estrictamente el esquema JSON proporcionado.",
    },
    "factura": {
        "schema_name": "factura_documento",
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "proveedor": {"type": "string"},
                "cliente": {"type": "string"},
                "numero_documento": {"type": "string"},
                "fecha_emision": {"type": "string"},
                "moneda": {"type": "string"},
                "subtotal": {"type": "string"},
                "impuestos": {"type": "string"},
                "total": {"type": "string"},
                "items": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "descripcion": {"type": "string"},
                            "cantidad": {"type": "string"},
                            "precio_unitario": {"type": "string"},
                            "total_linea": {"type": "string"},
                        },
                        "required": ["descripcion", "cantidad", "precio_unitario", "total_linea"],
                    },
                },
            },
            "required": [
                "proveedor",
                "cliente",
                "numero_documento",
                "fecha_emision",
                "moneda",
                "subtotal",
                "impuestos",
                "total",
                "items",
            ],
        },
        "prompt": "Extrae los campos de factura con precision. Mantiene importes y fechas exactamente como aparecen en el documento.",
    },
    "identidad": {
        "schema_name": "documento_identidad",
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "tipo_documento": {"type": "string"},
                "numero_documento": {"type": "string"},
                "nombres": {"type": "string"},
                "apellidos": {"type": "string"},
                "fecha_nacimiento": {"type": "string"},
                "fecha_emision": {"type": "string"},
                "fecha_vencimiento": {"type": "string"},
                "direccion": {"type": "string"},
            },
            "required": [
                "tipo_documento",
                "numero_documento",
                "nombres",
                "apellidos",
                "fecha_nacimiento",
                "fecha_emision",
                "fecha_vencimiento",
                "direccion",
            ],
        },
        "prompt": "Extrae campos de identidad de forma literal. Si no existe un dato, devuelve cadena vacia.",
    },
}


def _cleanup_file(path: str) -> None:
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass


def _safe_output_name(filename: str) -> str:
    name = Path(filename).stem or "documento"
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("._-")
    return normalized or "documento"


def _resolve_content_type(filename: str, declared: str) -> str:
    value = (declared or "").split(";")[0].strip().lower()
    if value in _SUPPORTED_CONTENT_TYPES:
        return value

    guessed, _ = mimetypes.guess_type(filename)
    guessed_value = (guessed or "").strip().lower()
    if guessed_value in _SUPPORTED_CONTENT_TYPES:
        return guessed_value

    if filename.lower().endswith(".pdf"):
        return "application/pdf"
    return "application/octet-stream"


def _validate_upload(file: UploadFile) -> None:
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    content_type = (file.content_type or "").split(";")[0].strip().lower()

    if ext in _SUPPORTED_EXTENSIONS:
        return
    if content_type in _SUPPORTED_CONTENT_TYPES:
        return

    raise HTTPException(
        status_code=400,
        detail="Formato no soportado. Sube un PDF o una imagen (PNG/JPG/TIFF/BMP/WEBP/HEIC)",
    )


def _write_txt_file(content: str) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp:
        tmp.write(content)
        return tmp.name


def _write_docx_file(content: str) -> str:
    doc = Document()
    lines = content.splitlines()
    if not lines:
        doc.add_paragraph(content)
    else:
        for line in lines:
            doc.add_paragraph(line)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = tmp.name
    doc.save(path)
    return path


def _write_json_file(payload: Dict[str, Any]) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".json", mode="w", encoding="utf-8") as tmp:
        json.dump(payload, tmp, ensure_ascii=False, indent=2)
        return tmp.name


async def _read_upload_bytes(file: UploadFile) -> Tuple[bytes, str, str]:
    """Shared upload validation+read for OCR endpoints."""
    _validate_upload(file)

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="El archivo esta vacio")

    max_bytes = max(1, int(settings.ocr_max_upload_mb)) * 1024 * 1024
    if len(raw) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"El archivo supera el limite de {settings.ocr_max_upload_mb} MB",
        )

    filename = os.path.basename(file.filename or "documento")
    content_type = _resolve_content_type(filename, file.content_type or "")
    return raw, filename, content_type


@router.get("/ocr-status")
async def ocr_status():
    """Returns the configured OCR backend and Ollama base URL."""
    return JSONResponse({
        "backend": settings.ocr_backend,
        "ollama_base_url": settings.ocr_ollama_base_url,
        "ollama_model_deepseek": settings.ocr_ollama_model_deepseek,
        "ollama_model_glm": settings.ocr_ollama_model_glm,
        "ollama_model_custom": settings.ocr_ollama_model_custom,
    })


@router.post("/ocr-probe")
async def probe_ollama(ollama_url: str = Form(...)):
    """Test connectivity to an Ollama server and list its available models."""
    url = (ollama_url or "").strip().rstrip("/")
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL debe comenzar con http:// o https://")
    if len(url) > _OLLAMA_URL_MAX_LENGTH:
        raise HTTPException(status_code=400, detail="URL demasiado larga")

    try:
        timeout = httpx.Timeout(timeout=8.0, connect=5.0)
        async with httpx.AsyncClient(timeout=timeout) as client:
            resp = await client.get(f"{url}/api/tags")
    except httpx.HTTPError as exc:
        return JSONResponse({"ok": False, "error": f"No se pudo conectar: {exc}"})

    if resp.status_code != 200:
        return JSONResponse({"ok": False, "error": f"Ollama respondió HTTP {resp.status_code}"})

    try:
        data = resp.json()
        models = [m["name"] for m in data.get("models", [])]
    except Exception:
        return JSONResponse({"ok": False, "error": "Respuesta inesperada de Ollama"})

    return JSONResponse({"ok": True, "models": models})


@router.post("/ocr-extract")
async def extract_ocr_to_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    output_format: Literal["txt", "docx"] = Form("txt"),
    ollama_url: str = Form(""),
    ollama_model: str = Form(""),
):
    raw, filename, content_type = await _read_upload_bytes(file)
    service = _resolve_service(ollama_url, ollama_model)

    try:
        result = await service.extract_text(
            file_bytes=raw,
            filename=filename,
            content_type=content_type,
        )
    except OCRConfigurationError as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except OCRServiceError as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Error interno OCR: {type(exc).__name__}: {exc}")

    if not result.text.strip():
        raise HTTPException(status_code=422, detail="No se detecto texto en el archivo")

    safe_base = _safe_output_name(filename)
    if output_format == "docx":
        output_path = _write_docx_file(result.text)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        download_name = f"{safe_base}_ocr.docx"
    else:
        output_path = _write_txt_file(result.text)
        media_type = "text/plain; charset=utf-8"
        download_name = f"{safe_base}_ocr.txt"

    background_tasks.add_task(_cleanup_file, output_path)
    return FileResponse(
        output_path,
        media_type=media_type,
        filename=download_name,
        headers={
            "X-OCR-Model": result.model,
            "X-OCR-Pages": str(result.pages_processed),
        },
    )


@router.post("/ocr-extract-structured")
async def extract_ocr_structured(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    schema_type: Literal["general", "factura", "identidad"] = Form("general"),
    instructions: str = Form(""),
    ollama_url: str = Form(""),
    ollama_model: str = Form(""),
):
    raw, filename, content_type = await _read_upload_bytes(file)
    service = _resolve_service(ollama_url, ollama_model)

    preset = _STRUCTURED_PRESETS[schema_type]
    extra_prompt = (instructions or "").strip()
    prompt = preset["prompt"]
    if extra_prompt:
        prompt = f"{prompt}\n\nInstrucciones adicionales del usuario:\n{extra_prompt}"

    try:
        result = await service.extract_structured(
            file_bytes=raw,
            filename=filename,
            content_type=content_type,
            schema_name=str(preset["schema_name"]),
            schema=dict(preset["schema"]),
            prompt=prompt,
        )
    except OCRConfigurationError as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except OCRServiceError as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Error interno OCR estructurado: {type(exc).__name__}: {exc}")

    if result.data in ({}, [], "", None):
        raise HTTPException(status_code=422, detail="No se pudo extraer estructura del documento")

    payload = {
        "schema_type": schema_type,
        "model": result.model,
        "pages_processed": result.pages_processed,
        "data": result.data,
    }
    safe_base = _safe_output_name(filename)
    output_path = _write_json_file(payload)

    background_tasks.add_task(_cleanup_file, output_path)
    return FileResponse(
        output_path,
        media_type="application/json; charset=utf-8",
        filename=f"{safe_base}_ocr_{schema_type}.json",
        headers={
            "X-OCR-Model": result.model,
            "X-OCR-Pages": str(result.pages_processed),
            "X-OCR-Schema": schema_type,
        },
    )
